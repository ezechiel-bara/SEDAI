import os
import sys
import subprocess

# 1. Vérifier et installer python-docx si absent
try:
    import docx
except ImportError:
    print("python-docx non trouvé. Installation en cours...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generate_word_report():
    # Chemins
    base_dir = r"c:\code\SEDAI"
    md_path = r"C:\Users\YankeMall\.gemini\antigravity\brain\04a11407-cab0-40ae-8a06-d1c1d2d4e22a\architecture_diagnostic_auto.md"
    docx_output_dir = os.path.join(base_dir, "docs")
    docx_path = os.path.join(docx_output_dir, "architecture_diagnostic_auto.docx")
    
    os.makedirs(docx_output_dir, exist_ok=True)
    
    if not os.path.exists(md_path):
        print(f"Erreur : Le fichier source Markdown {md_path} n'existe pas.")
        return False
        
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    doc = Document()
    
    # Configuration des marges
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles par défaut
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2B, 0x2D, 0x2F) # Anthracite chic
    
    print("Génération du document Word...")
    
    in_table = False
    table_data = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Ignorer les balises Mermaid pour le fichier Word (non prises en charge directement)
        if line.startswith("```mermaid"):
            while i < len(lines) and not lines[i].strip() == "```":
                i += 1
            i += 1
            continue
            
        # Détection des tableaux Markdown
        if line.startswith("|"):
            if not in_table:
                in_table = True
                table_data = []
            
            # Ne pas ajouter la ligne de séparation (ex: | :--- | :--- |)
            if not ("---" in line):
                # Nettoyer et couper
                cells = [c.strip() for c in line.split("|")[1:-1]]
                table_data.append(cells)
            i += 1
            continue
        else:
            if in_table:
                # Écrire le tableau accumulé
                if table_data:
                    headers = table_data[0]
                    rows = table_data[1:]
                    
                    word_table = doc.add_table(rows=len(table_data), cols=len(headers))
                    word_table.style = 'Light Shading Accent 1'
                    
                    # Remplir l'en-tête
                    for col_idx, text in enumerate(headers):
                        cell = word_table.cell(0, col_idx)
                        cell.text = text
                        set_cell_background(cell, "1F497D")  # Bleu marine premium
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.name = 'Segoe UI'
                                
                    # Remplir le contenu
                    for row_idx, row_cells in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_cells):
                            cell = word_table.cell(row_idx + 1, col_idx)
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Segoe UI'
                                    run.font.size = Pt(10)
                                    
                in_table = False
                doc.add_paragraph() # Espacement après tableau
                
        # Titre Principal
        if line.startswith("# "):
            title_text = line[2:].replace("🛠️ ", "")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.font.name = 'Segoe UI Semibold'
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            p.paragraph_format.space_after = Pt(24)
            
        # Titre Niveau 2 (##)
        elif line.startswith("## "):
            section_text = line[3:]
            p = doc.add_paragraph()
            run = p.add_run(section_text)
            run.font.name = 'Segoe UI Semibold'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            
        # Titre Niveau 3 (###)
        elif line.startswith("### "):
            sub_text = line[4:]
            p = doc.add_paragraph()
            run = p.add_run(sub_text)
            run.font.name = 'Segoe UI'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
        # Alertes ou Citations (> [!NOTE] etc)
        elif line.startswith(">"):
            alert_text = line[1:].strip()
            # Si c'est l'indicateur d'alerte, lire le texte réel en dessous ou à côté
            if alert_text.startswith("[!"):
                i += 1
                if i < len(lines):
                    alert_text = lines[i].strip("> \n")
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            
            # Ajouter une bordure gauche ou un style de panneau
            run = p.add_run("ℹ️ Note : " + alert_text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            
        # Liste à puces (* ou -)
        elif line.startswith("*") or line.startswith("-"):
            bullet_text = line[1:].strip()
            
            # Formater le gras inline (ex: **Texte**)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            
            # Parsing gras basique
            parts = bullet_text.split("**")
            is_bold = False
            for part in parts:
                run = p.add_run(part)
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                if is_bold:
                    run.font.bold = True
                is_bold = not is_bold
                
        # Lignes horizontales
        elif line.startswith("---"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="D3D3D3"/></w:pBdr>')
            p._p.get_or_add_pPr().append(p_border)
            
        # Paragraphes normaux
        elif line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            
            # Parsing gras basique
            parts = line.split("**")
            is_bold = False
            for part in parts:
                # Éviter de mettre les formules LaTeX brutes en format bizarre
                run = p.add_run(part)
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                if is_bold:
                    run.font.bold = True
                is_bold = not is_bold
                
        i += 1
        
    doc.save(docx_path)
    print(f"Document Word enregistré avec succès dans : {docx_path}")
    return True

if __name__ == "__main__":
    generate_word_report()
